# -*- coding: utf-8 -*-
"""Smoke tests for LeAgent-ported docread tools."""
from __future__ import annotations

import tempfile
from pathlib import Path

import pytest


@pytest.mark.asyncio
async def test_csv_excel_word_pdf_html_md_text_smoke():
    from openpyxl import Workbook
    from docx import Document
    import fitz

    from reactor_tool.tool.docread.service import (
        run_csv_processor,
        run_excel_reader,
        run_word_reader,
        run_pdf_reader,
        run_pdf_structure,
        run_citation_extractor,
        run_html_processor,
        run_markdown_processor,
        run_text_processor,
    )

    tmp = Path(tempfile.mkdtemp())
    csv_path = tmp / "a.csv"
    csv_path.write_text("name,age\nalice,1\nbob,2\n", encoding="utf-8")

    xlsx = tmp / "a.xlsx"
    wb = Workbook()
    wb.active["A1"] = "name"
    wb.active["B1"] = "age"
    wb.active["A2"] = "alice"
    wb.active["B2"] = 1
    wb.save(xlsx)

    docx_path = tmp / "a.docx"
    d = Document()
    d.add_heading("Title", 1)
    d.add_paragraph("para")
    d.save(docx_path)

    pdf_path = tmp / "a.pdf"
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((72, 72), "Hello PDF Abstract Introduction References [1] Paper.")
    doc.save(pdf_path)
    doc.close()

    html_path = tmp / "a.html"
    html_path.write_text(
        "<html><head><title>T</title></head><body><p>Hi</p></body></html>",
        encoding="utf-8",
    )
    md_path = tmp / "a.md"
    md_path.write_text("# Hello\n\nworld\n", encoding="utf-8")
    txt_path = tmp / "a.txt"
    txt_path.write_text("hello world\n", encoding="utf-8")

    rid = "test-docread-smoke"
    assert (await run_csv_processor(rid, {"operation": "read", "file_path": str(csv_path)}))["success"]
    excel = await run_excel_reader(rid, {"file_path": str(xlsx)})
    assert excel["success"] and excel["data"]["row_count"] == 1
    word = await run_word_reader(rid, {"file_path": str(docx_path)})
    assert word["success"] and word["data"]["heading_count"] == 1
    pdf = await run_pdf_reader(rid, {"operation": "read", "file_path": str(pdf_path)})
    assert pdf["success"]
    struct = await run_pdf_structure(rid, {"file_path": str(pdf_path)})
    assert struct["success"] and struct["data"]["page_count"] == 1
    cites = await run_citation_extractor(rid, {"file_path": str(pdf_path)})
    assert cites["success"]
    assert (await run_html_processor(rid, {"operation": "read", "file_path": str(html_path)}))["success"]
    assert (await run_markdown_processor(rid, {"operation": "read", "file_path": str(md_path)}))["success"]
    assert (await run_text_processor(rid, {"operation": "read", "file_path": str(txt_path)}))["success"]


def test_image_ocr_validates_extension_and_empty():
    from reactor_tool.tool.docread.image_ocr import run_image_ocr_sync

    tmp = Path(tempfile.mkdtemp())
    bad = tmp / "a.txt"
    bad.write_text("not image", encoding="utf-8")
    try:
        run_image_ocr_sync(str(bad))
        assert False, "expected ValueError"
    except ValueError as e:
        assert "Unsupported image type" in str(e)

    empty = tmp / "empty.png"
    empty.write_bytes(b"")
    try:
        run_image_ocr_sync(str(empty))
        assert False, "expected ValueError"
    except ValueError as e:
        assert "empty" in str(e).lower()


def test_docread_finds_persistent_file_store_after_release_switch(monkeypatch, tmp_path):
    from reactor_tool.tool.docread.paths import resolve_input_path

    persistent_root = tmp_path / "data" / "file_db_dir"
    session_root = persistent_root / "session-release-switch"
    session_root.mkdir(parents=True)
    source = session_root / "logs.csv"
    source.write_text("name\nvalue\n", encoding="utf-8")

    monkeypatch.setenv("FILE_SAVE_PATH", str(persistent_root))
    (tmp_path / "new-release").mkdir()
    monkeypatch.chdir(tmp_path / "new-release")

    assert resolve_input_path(
        "logs.csv",
        request_id="session-release-switch",
    ) == source.resolve()


def test_html_processor_accepts_legacy_metadata_operation_name():
    import asyncio

    from reactor_tool.tool.docread.service import run_html_processor

    tmp = Path(tempfile.mkdtemp())
    html_path = tmp / "metadata.html"
    html_path.write_text(
        '<html><head><title>Test page</title><meta name="description" content="demo"></head></html>',
        encoding="utf-8",
    )

    result = asyncio.run(
        run_html_processor(
            "test-html-metadata-alias",
            {"operation": "metadata", "file_path": str(html_path)},
        )
    )

    assert result["success"]
    assert result["data"]["title"] == "Test page"


def test_image_ocr_uses_mock_backend(monkeypatch):
    from reactor_tool.tool import docread
    from reactor_tool.tool.docread import image_ocr as image_ocr_mod

    tmp = Path(tempfile.mkdtemp())
    img = tmp / "shot.png"
    # minimal valid-enough file for size checks (not decoded by OCR when mocked)
    img.write_bytes(b"\x89PNG\r\n\x1a\n" + b"\x00" * 64)

    class _Fake:
        def ocr(self, path: str) -> str:
            assert path.endswith("shot.png")
            return "  hello from cloud ocr  "

    import reactor_tool.tool.mrag.utils.ocr_utils as ocr_utils

    monkeypatch.setattr(ocr_utils, "get_ocr_model", lambda: _Fake())
    monkeypatch.setenv("OCR_TYPE", "vlm-ocr")
    data = image_ocr_mod.run_image_ocr_sync(str(img))
    assert data["text"] == "hello from cloud ocr"
    assert data["char_count"] == len("hello from cloud ocr")
    assert data["ocr_type"] == "vlm-ocr"
